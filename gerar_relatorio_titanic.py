import pandas as pd
import seaborn as sns
import matplotlib.pyplot as plt
import os
from datetime import datetime
from docx import Document
from docx.shared import Inches
from fpdf import FPDF
from sklearn.ensemble import RandomForestClassifier
from sklearn.model_selection import train_test_split
from sklearn.metrics import classification_report, confusion_matrix, accuracy_score
from sklearn.preprocessing import LabelEncoder
from scipy.stats import chi2_contingency

# Configuração de Estilo dos Gráficos
sns.set_theme(style="whitegrid")
plt.rcParams.update({'figure.max_open_warning': 0})

class RelatorioBuilder:
    """
    Classe responsável por orquestrar a geração de conteúdo idêntico
    para Markdown, DOCX e PDF.
    """
    def __init__(self, titulo):
        self.titulo = titulo
        self.conteudo = [] # Lista de dicionários {'tipo': 'texto|img|tabela|titulo', 'valor': ...}

    def adicionar_titulo(self, texto, nivel=1):
        self.conteudo.append({'tipo': 'titulo', 'valor': texto, 'nivel': nivel})

    def adicionar_texto(self, texto):
        self.conteudo.append({'tipo': 'texto', 'valor': texto})

    def adicionar_imagem(self, caminho_img, legenda=""):
        self.conteudo.append({'tipo': 'imagem', 'valor': caminho_img, 'legenda': legenda})

    def adicionar_tabela(self, dataframe, legenda=""):
        self.conteudo.append({'tipo': 'tabela', 'valor': dataframe, 'legenda': legenda})

    def salvar_md(self, nome_arquivo):
        with open(nome_arquivo, 'w', encoding='utf-8') as f:
            f.write(f"# {self.titulo}\n\n")
            for item in self.conteudo:
                if item['tipo'] == 'titulo':
                    f.write(f"{'#' * (item['nivel'] + 1)} {item['valor']}\n\n")
                elif item['tipo'] == 'texto':
                    f.write(f"{item['valor']}\n\n")
                elif item['tipo'] == 'imagem':
                    f.write(f"![{item['legenda']}]({item['valor']})\n")
                    f.write(f"*{item['legenda']}*\n\n")
                elif item['tipo'] == 'tabela':
                    f.write(f"### {item['legenda']}\n")
                    f.write(item['valor'].to_markdown(index=False))
                    f.write("\n\n")
        print(f"[SUCESSO] Relatório Markdown salvo em: {nome_arquivo}")

    def salvar_docx(self, nome_arquivo):
        doc = Document()
        doc.add_heading(self.titulo, 0)

        for item in self.conteudo:
            if item['tipo'] == 'titulo':
                doc.add_heading(item['valor'], level=item['nivel'])
            elif item['tipo'] == 'texto':
                doc.add_paragraph(item['valor'])
            elif item['tipo'] == 'imagem':
                try:
                    doc.add_picture(item['valor'], width=Inches(5.5))
                    doc.add_paragraph(item['legenda'], style='Caption')
                except Exception as e:
                    doc.add_paragraph(f"[Erro ao inserir imagem: {e}]")
            elif item['tipo'] == 'tabela':
                doc.add_paragraph(item['legenda'], style='Caption')
                df = item['valor']
                t = doc.add_table(rows=df.shape[0] + 1, cols=df.shape[1])
                t.style = 'Table Grid'
                # Header
                for j, col_name in enumerate(df.columns):
                    t.cell(0, j).text = str(col_name)
                # Body
                for i, row in enumerate(df.itertuples(index=False)):
                    for j, val in enumerate(row):
                        t.cell(i + 1, j).text = str(val)
                doc.add_paragraph("") # Espaço

        doc.save(nome_arquivo)
        print(f"[SUCESSO] Relatório DOCX salvo em: {nome_arquivo}")

    def salvar_pdf(self, nome_arquivo):
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()
        
        # Configuração básica de fonte (Arial é padrão no FPDF)
        pdf.set_font("Arial", "B", 16)
        pdf.cell(0, 10, self.titulo.encode('latin-1', 'replace').decode('latin-1'), ln=True, align='C')
        pdf.ln(10)

        for item in self.conteudo:
            texto_safe = str(item['valor']).encode('latin-1', 'replace').decode('latin-1')
            
            if item['tipo'] == 'titulo':
                pdf.set_font("Arial", "B", 14 - item['nivel'])
                pdf.cell(0, 10, texto_safe, ln=True)
                pdf.ln(2)
            
            elif item['tipo'] == 'texto':
                pdf.set_font("Arial", "", 11)
                pdf.multi_cell(0, 6, texto_safe)
                pdf.ln(5)
            
            elif item['tipo'] == 'imagem':
                try:
                    # Centralizar imagem
                    largura_pag = pdf.w - 2 * pdf.l_margin
                    pdf.image(item['valor'], x=pdf.l_margin, w=largura_pag)
                    pdf.set_font("Arial", "I", 9)
                    legenda_safe = item['legenda'].encode('latin-1', 'replace').decode('latin-1')
                    pdf.cell(0, 5, legenda_safe, ln=True, align='C')
                    pdf.ln(5)
                except:
                    pdf.cell(0, 10, "[Imagem nao encontrada]", ln=True)

            elif item['tipo'] == 'tabela':
                pdf.set_font("Arial", "B", 10)
                legenda_safe = item['legenda'].encode('latin-1', 'replace').decode('latin-1')
                pdf.cell(0, 8, legenda_safe, ln=True)
                
                df = item['valor']
                # Largura simples das colunas
                col_width = (pdf.w - 2 * pdf.l_margin) / len(df.columns)
                pdf.set_font("Arial", "B", 9)
                
                # Header
                for col in df.columns:
                    col_safe = str(col).encode('latin-1', 'replace').decode('latin-1')
                    pdf.cell(col_width, 7, col_safe, border=1)
                pdf.ln()
                
                # Body
                pdf.set_font("Arial", "", 9)
                for row in df.itertuples(index=False):
                    for val in row:
                        val_safe = str(val).encode('latin-1', 'replace').decode('latin-1')
                        pdf.cell(col_width, 6, val_safe, border=1)
                    pdf.ln()
                pdf.ln(5)

        pdf.output(nome_arquivo)
        print(f"[SUCESSO] Relatório PDF salvo em: {nome_arquivo}")

def analisar_dados_titanic():
    # 1. Carregamento e Limpeza
    try:
        df = pd.read_csv('train.csv')
    except FileNotFoundError:
        # Criação de dados dummy caso o arquivo não exista para o exemplo rodar
        print("AVISO: 'train.csv' não encontrado. Gerando dados de exemplo.")
        data = {
            'Survived': [0, 1, 1, 0, 0, 1, 0, 1, 1, 0] * 10,
            'Pclass': [3, 1, 3, 1, 3, 2, 3, 1, 2, 3] * 10,
            'Sex': ['male', 'female', 'female', 'female', 'male', 'female', 'male', 'female', 'female', 'male'] * 10,
            'Age': [22, 38, 26, 35, 35, 27, 2, 14, 4, 50] * 10,
            'Fare': [7.25, 71.28, 7.92, 53.1, 8.05, 11.13, 21.07, 30.07, 16.7, 8.05] * 10,
            'Embarked': ['S', 'C', 'S', 'S', 'S', 'Q', 'S', 'S', 'S', 'S'] * 10,
             'SibSp': [1, 1, 0, 1, 0, 0, 0, 3, 0, 1] * 10,
            'Parch': [0, 0, 0, 0, 0, 0, 0, 1, 2, 0] * 10,
            'Name': ['Braund, Mr. Owen', 'Cumings, Mrs. John', 'Heikkinen, Miss. Laina', 'Futrelle, Mrs. Jacques', 'Allen, Mr. William', 'Moran, Mr. James', 'McCarthy, Mr. Timothy', 'Palsson, Master. Gosta', 'Johnson, Mrs. Oscar', 'Nasser, Mrs. Nicholas'] * 10,
        }
        df = pd.DataFrame(data)

    # Tratamento básico
    df['Age'] = df['Age'].fillna(df['Age'].median())
    df['Embarked'] = df['Embarked'].fillna(df['Embarked'].mode()[0])
    
    # 2. Geração de Gráficos
    os.makedirs('output/graficos', exist_ok=True)
    
    # Gráfico 1: Sobrevivência por Sexo
    plt.figure(figsize=(8, 5))
    sns.countplot(data=df, x='Sex', hue='Survived', palette='viridis')
    plt.title('Sobrevivência por Sexo')
    plt.xlabel('Sexo')
    plt.ylabel('Contagem')
    plt.legend(title='Sobreviveu', labels=['Não', 'Sim'])
    img_sexo = 'output/graficos/sobrevivencia_sexo.png'
    plt.savefig(img_sexo)
    plt.close()

    # Gráfico 2: Sobrevivência por Classe
    plt.figure(figsize=(8, 5))
    sns.barplot(data=df, x='Pclass', y='Survived', hue='Pclass', palette='magma', errorbar=None, legend=False)
    plt.title('Taxa de Sobrevivência por Classe')
    plt.xlabel('Classe (1 = Alta, 3 = Baixa)')
    plt.ylabel('Taxa de Sobrevivência (0 a 1)')
    img_classe = 'output/graficos/sobrevivencia_classe.png'
    plt.savefig(img_classe)
    plt.close()

    # Gráfico 3: Distribuição de Idade
    plt.figure(figsize=(8, 5))
    sns.histplot(data=df, x='Age', hue='Survived', kde=True, element="step", palette='coolwarm')
    plt.title('Distribuição de Idade e Sobrevivência')
    img_idade = 'output/graficos/distribuicao_idade.png'
    plt.savefig(img_idade)
    plt.close()

    # Gráfico 4: Correlação (Heatmap)
    plt.figure(figsize=(10, 8))
    numeric_df = df.select_dtypes(include=['float64', 'int64'])
    corr = numeric_df.corr()
    sns.heatmap(corr, annot=True, cmap='coolwarm', fmt=".2f")
    plt.title('Mapa de Calor de Correlação')
     img_corr = 'output/graficos/correlacao.png'
    plt.savefig(img_corr)
    plt.close()

    # Gráfico 9: Teste Qui-Quadrado (Heatmap Categórico)
    categorical_cols = ['Survived', 'Pclass', 'Sex', 'Embarked']
    chi2_matrix = pd.DataFrame(index=categorical_cols, columns=categorical_cols, dtype=float)
    for col1 in categorical_cols:
        for col2 in categorical_cols:
            if col1 == col2:
                chi2_matrix.loc[col1, col2] = 0.0
            else:
                contingency = pd.crosstab(df[col1], df[col2])
                chi2, p, _, _ = chi2_contingency(contingency)
                chi2_matrix.loc[col1, col2] = p
    plt.figure(figsize=(8, 6))
    sns.heatmap(chi2_matrix, annot=True, cmap='coolwarm_r', fmt=".2e")
    plt.title('P-valores do Teste Qui-Quadrado (Associação)')
    img_chi2 = 'output/graficos/chi2_heatmap.png'
    plt.savefig(img_chi2)
    plt.close()

    # 3. Modelagem Simples (Machine Learning)
    le = LabelEncoder()
    df_ml = df.copy()
    df_ml['Sex'] = le.fit_transform(df_ml['Sex'])
    df_ml['Embarked'] = le.fit_transform(df_ml['Embarked'])
    X = df_ml[['Pclass', 'Sex', 'Age', 'Fare']]
    y = df_ml['Survived']
    
    X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.3, random_state=42)
    
    model = RandomForestClassifier(n_estimators=100, random_state=42)
    model.fit(X_train, y_train)
    y_pred = model.predict(X_test)
    
    acc = accuracy_score(y_test, y_pred)
    report = classification_report(y_test, y_pred, output_dict=True)
    
    # Importância das Features
    feature_imp = pd.DataFrame({
        'Feature': X.columns,
        'Importancia': model.feature_importances_
    }).sort_values(by='Importancia', ascending=False)

    return {
        'df_head': df.head(),
        'stats': df.describe(),
        'imgs': {'sexo': img_sexo, 'classe': img_classe, 'idade': img_idade, 'cm': img_cm, 'corr': img_corr, 'family': img_family, 'chi2': img_chi2},
        'ml': {'acc': acc, 'report': report, 'feature_imp': feature_imp},
        'counts': df['Survived'].value_counts(normalize=True)
    }

def gerar_relatorio_completo():
    print("Iniciando análise de dados...")
    dados = analisar_dados_titanic()

    builder = RelatorioBuilder("Relatório Final: Análise do Titanic")

    # --- INTRODUÇÃO ---
    builder.adicionar_titulo("1. Introdução e Contexto", 1)
    builder.adicionar_texto(
        "Este relatório apresenta uma análise detalhada sobre o naufrágio do RMS Titanic. "
        "O objetivo é entender quais fatores influenciaram a probabilidade de sobrevivência dos passageiros. "
        "Utilizamos técnicas de Ciência de Dados e Aprendizado de Máquina para extrair padrões dos dados históricos."
    )
    builder.adicionar_texto(
        "**Para o leitor leigo:** Imagine que estamos tentando descobrir se a sorte, o dinheiro ou a idade "
        "foram mais importantes para salvar uma vida naquela noite. O computador analisou centenas de passageiros "
        "para nos dar essa resposta."
    )
    builder.adicionar_texto(
        "**Para o especialista:** O dataset passou por pré-processamento (imputação de nulos na idade pela mediana "
        "e moda para embarque). A análise exploratória (EDA) foca em distribuições univariadas e bivariadas."
    )

    # --- DADOS ---
    builder.adicionar_titulo("2. Visão Geral dos Dados", 1)
    builder.adicionar_texto(
        "Abaixo, uma amostra das primeiras linhas do conjunto de dados utilizado. "
        "Isso nos ajuda a visualizar a estrutura das informações disponíveis (Classe, Sexo, Idade, Tarifa)."
    )
    builder.adicionar_tabela(dados['df_head'], "Amostra dos Dados (Primeiras 5 linhas)")
    
    builder.adicionar_titulo("Estatísticas Descritivas", 2)
    builder.adicionar_texto(
        "A tabela a seguir resume matematicamente os dados numéricos (média, desvio padrão, mínimo, máximo)."
    )
    # Simplificando a tabela de stats para caber melhor no PDF/DOCX
    stats_simple = dados['stats'].loc[['mean', 'min', 'max', 'std']].round(2).reset_index()
    builder.adicionar_tabela(stats_simple, "Resumo Estatístico")

    # --- ANÁLISE VISUAL ---
    builder.adicionar_titulo("3. Análise Exploratória Visual", 1)
    
    # Sexo
    builder.adicionar_titulo("3.1 Influência do Gênero", 2)
    builder.adicionar_texto(
        "Historicamente, a regra 'mulheres e crianças primeiro' foi aplicada. "
        "O gráfico abaixo confirma se os dados refletem essa regra."
    )
    builder.adicionar_imagem(dados['imgs']['sexo'], "Comparação de Sobrevivência entre Homens e Mulheres")
    builder.adicionar_texto(
        "**Interpretação:** Observa-se uma taxa de sobrevivência significativamente maior para o sexo feminino. "
        "Isso indica uma forte correlação entre gênero e sobrevivência."
    )

    # Classe
    builder.adicionar_titulo("3.2 Influência da Classe Social (Poder Econômico)", 2)
    builder.adicionar_texto(
        "A classe do passageiro (1ª, 2ª ou 3ª) é um indicativo socioeconômico. "
        "Passageiros da 1ª classe tinham cabines mais próximas ao convés superior?"
    )
    builder.adicionar_imagem(dados['imgs']['classe'], "Taxa de Sobrevivência por Classe")
    builder.adicionar_texto(
        "**Interpretação:** Passageiros da 1ª classe tiveram as maiores chances de sobrevivência, "
        "enquanto a 3ª classe sofreu as maiores perdas proporcionais."
    )

    # Idade
    builder.adicionar_titulo("3.3 Distribuição de Idade", 2)
    builder.adicionar_texto(
        "Analisamos como a idade impactou as chances. Crianças foram salvas? Idosos tiveram prioridade?"
    )
    builder.adicionar_imagem(dados['imgs']['idade'], "Histograma de Idade: Sobreviventes vs Não Sobreviventes")
    builder.adicionar_texto(
        "**Interpretação:** Nota-se um pico de sobrevivência em crianças pequenas (0-5 anos), "
        "validando a prioridade dada aos infantes."
    )

    # Correlação
    builder.adicionar_titulo("3.4 Análise de Correlação (Heatmap)", 2)
    builder.adicionar_texto(
        "O mapa de calor abaixo exibe a correlação entre as variáveis numéricas. "
        "Correlação mede a relação estatística entre duas variáveis."
    )
    builder.adicionar_imagem(dados['imgs']['corr'], "Matriz de Correlação das Variáveis")
    builder.adicionar_texto(
        "**Interpretação:**\n"
        "- **Cores Quentes (Vermelho):** Correlação positiva (ambas sobem juntas).\n"
        "- **Cores Frias (Azul):** Correlação negativa (uma sobe, a outra desce).\n"
        "- **Destaque:** Existe uma correlação negativa relevante entre Classe (Pclass) e Tarifa (Fare)."
    )

    # --- MACHINE LEARNING ---
    builder.adicionar_titulo("4. Modelagem Preditiva (Machine Learning)", 1)
    builder.adicionar_texto(
        "Utilizamos um algoritmo chamado **Random Forest** (Floresta Aleatória). "
        "Ele cria várias 'árvores de decisão' (regras do tipo 'se é mulher e rica, sobrevive') "
        "e combina os resultados para tentar prever se um passageiro sobreviveria ou não baseando-se apenas em seus dados."
    )
    
    acc_percent = dados['ml']['acc'] * 100
    builder.adicionar_texto(
        f"**Acurácia do Modelo:** O modelo acertou **{acc_percent:.2f}%** das previsões nos dados de teste."
    )

    builder.adicionar_titulo("4.1 O que foi mais importante para o modelo?", 2)
    builder.adicionar_texto(
        "O algoritmo nos diz quais características pesaram mais na decisão de classificar alguém como sobrevivente."
    )
    builder.adicionar_tabela(dados['ml']['feature_imp'].round(4), "Importância das Variáveis (Feature Importance)")
    
    builder.adicionar_texto(
        "**Análise Técnica:** A 'Feature Importance' derivada do Random Forest (baseada na impureza de Gini) "
        "mostra que Sexo, Idade e Tarifa são preditores cruciais. Isso corrobora a análise visual feita anteriormente."
    )

    builder.adicionar_titulo("4.3 Análise de Correlação (Teste Qui-Quadrado)", 2)
    builder.adicionar_texto(
        "A matriz abaixo exibe a correlação entre as variáveis categóricas. "
        "Correlação mede a relação estatística entre duas variáveis."
    )
    builder.adicionar_imagem(dados['imgs']['chi2'], "Matriz de Correlação das Variáveis Categóricas")
    builder.adicionar_texto(
        "**Interpretação:** \n"
        "Os valores de cada célula representam o nível de correlação entre as duas variáveis, quanto mais próximo de zero maior a correlação."
    )
    # --- CONCLUSÃO ---
    builder.adicionar_titulo("5. Conclusão", 1)
    builder.adicionar_texto(
        "A análise dos dados do Titanic revela que a sobrevivência não foi aleatória. "
        "Protocolos de emergência e status social desempenharam papéis fundamentais."
    )
    builder.adicionar_texto(
        "**Pontos Chave:**\n"
        "1. **Mulheres** tiveram prioridade absoluta.\n"
        "2. **Crianças** pequenas foram protegidas.\n"
        "3. **Classe Social** importou: dinheiro comprou segurança (ou melhor acesso aos botes).\n"
    )
    builder.adicionar_texto(
        "Este relatório demonstra como dados históricos podem ser transformados em conhecimento "
        "acionável e compreensível tanto para o público geral quanto para a academia."
    )
    
    data_atual = datetime.now().strftime("%d/%m/%Y")
    builder.adicionar_texto(f"\nRelatório gerado automaticamente em: {data_atual}")

    # --- SALVAR ARQUIVOS ---
    print("Gerando arquivos finais...")
    builder.salvar_md("Relatorio_Final_Titanic.md")
    builder.salvar_docx("Relatorio_Final_Titanic.docx")
    builder.salvar_pdf("Relatorio_Final_Titanic.pdf")
    print("Processo concluído com sucesso!")

if __name__ == "__main__":
    gerar_relatorio_completo()