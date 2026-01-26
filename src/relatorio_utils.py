import subprocess
import sys

def check_dependencies(dependencies):
    missing = []
    for dependency in dependencies:
        try:
            __import__(dependency)
        except ImportError:
            missing.append(dependency)
    if missing:
        raise ImportError(f"As seguintes dependências estão faltando: {', '.join(missing)}")

import os
from docx import Document
from docx.shared import Inches
from fpdf import FPDF
import pandas as pd

# Verificação de dependências do módulo
dependencies = ['docx', 'fpdf', 'pandas']
check_dependencies(dependencies)

class RelatorioBuilder:
    """
    Classe responsável por orquestrar a geração de conteúdo idêntico
    para Markdown, DOCX e PDF.
    """
    def __init__(self, titulo):
        self.titulo = titulo
        self.conteudo = [] # Lista de dicionários {'tipo': 'texto|img|tabela|titulo', 'valor': ...}

    def adicionar_titulo(self, texto, nivel=1):
        self.conteudo.append({'tipo': 'titulo', 'valor': texto, 'nivel': nivel})

    def adicionar_texto(self, texto):
        self.conteudo.append({'tipo': 'texto', 'valor': texto})

    def adicionar_imagem(self, caminho_img, legenda=""):
        self.conteudo.append({'tipo': 'imagem', 'valor': caminho_img, 'legenda': legenda})

    def adicionar_tabela(self, dataframe, legenda=""):
        self.conteudo.append({'tipo': 'tabela', 'valor': dataframe, 'legenda': legenda})

    def salvar_md(self, nome_arquivo):
        with open(nome_arquivo, 'w', encoding='utf-8') as f:
            f.write(f"# {self.titulo}\n\n")
            for item in self.conteudo:
                if item['tipo'] == 'titulo':
                    f.write(f"{'#' * (item['nivel'] + 1)} {item['valor']}\n\n")
                elif item['tipo'] == 'texto':
                    f.write(f"{item['valor']}\n\n")
                elif item['tipo'] == 'imagem':
                    f.write(f"![{item['legenda']}]({item['valor']})\n")
                    f.write(f"*{item['legenda']}*\n\n")
                elif item['tipo'] == 'tabela':
                    f.write(f"### {item['legenda']}\n")
                    f.write(item['valor'].to_markdown(index=False))
                    f.write("\n\n")
        print(f"[SUCESSO] Relatório Markdown salvo em: {nome_arquivo}")

    def salvar_docx(self, nome_arquivo):
        doc = Document()
        doc.add_heading(self.titulo, 0)

        for item in self.conteudo:
            if item['tipo'] == 'titulo':
                doc.add_heading(item['valor'], level=item['nivel'])
            elif item['tipo'] == 'texto':
                doc.add_paragraph(item['valor'])
            elif item['tipo'] == 'imagem':
                try:
                    doc.add_picture(item['valor'], width=Inches(5.5))
                    doc.add_paragraph(item['legenda'], style='Caption')
                except Exception as e:
                    doc.add_paragraph(f"[Erro ao inserir imagem: {e}]")
            elif item['tipo'] == 'tabela':
                doc.add_paragraph(item['legenda'], style='Caption')
                df = item['valor']
                t = doc.add_table(rows=df.shape[0] + 1, cols=df.shape[1])
                t.style = 'Table Grid'
                # Header
                for j, col_name in enumerate(df.columns):
                    t.cell(0, j).text = str(col_name)
                # Body
                for i, row in enumerate(df.itertuples(index=False)):
                    for j, val in enumerate(row):
                        t.cell(i + 1, j).text = str(val)
                doc.add_paragraph("") # Espaço

        doc.save(nome_arquivo)
        print(f"[SUCESSO] Relatório DOCX salvo em: {nome_arquivo}")

    def salvar_pdf(self, nome_arquivo):
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()
        
        # Configuração básica de fonte (Arial é padrão no FPDF)
        pdf.set_font("Arial", "B", 16)
        pdf.cell(0, 10, self.titulo.encode('latin-1', 'replace').decode('latin-1'), ln=True, align='C')
        pdf.ln(10)

        for item in self.conteudo:
            texto_safe = str(item['valor']).encode('latin-1', 'replace').decode('latin-1')
            
            if item['tipo'] == 'titulo':
                pdf.set_font("Arial", "B", 14 - item['nivel'])
                pdf.cell(0, 10, texto_safe, ln=True)
                pdf.ln(2)
            
            elif item['tipo'] == 'texto':
                pdf.set_font("Arial", "", 11)
                pdf.multi_cell(0, 6, texto_safe)
                pdf.ln(5)
            
            elif item['tipo'] == 'imagem':
                try:
                    # Centralizar imagem
                    largura_pag = pdf.w - 2 * pdf.l_margin
                    pdf.image(item['valor'], x=pdf.l_margin, w=largura_pag)
                    pdf.set_font("Arial", "I", 9)
                    legenda_safe = item['legenda'].encode('latin-1', 'replace').decode('latin-1')
                    pdf.cell(0, 5, legenda_safe, ln=True, align='C')
                    pdf.ln(5)
                except:
                    pdf.cell(0, 10, "[Imagem nao encontrada]", ln=True)

            elif item['tipo'] == 'tabela':
                # Implementação simplificada de tabela para PDF
                pdf.set_font("Arial", "B", 10)
                legenda_safe = item['legenda'].encode('latin-1', 'replace').decode('latin-1')
                pdf.cell(0, 8, legenda_safe, ln=True)
                pdf.ln(2)

        pdf.output(nome_arquivo)
        print(f"[SUCESSO] Relatório PDF salvo em: {nome_arquivo}")
